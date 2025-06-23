import streamlit as st

# Configure the page at the very start - must be first Streamlit command
st.set_page_config(
    page_title="RAG Chat",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Debug mode UI at the top
st.title("RAG Chat Debug Mode")
st.info("If you see this, Streamlit is working. If the app fails below, you'll see an error message here.")

try:
    import os
    os.environ.pop("HTTP_PROXY", None)
    os.environ.pop("HTTPS_PROXY", None)
    
    import tempfile
    import time
    import sys
    import json
    import pandas as pd
    import uuid
    from typing import List, Dict, Any, Optional, Tuple
    import traceback

    # Fix protobuf compatibility issues with ChromaDB
    os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

    # Import required libraries
    import cohere
    from openai import OpenAI
    from dotenv import load_dotenv
    from langchain_community.document_loaders import PyPDFLoader, WebBaseLoader, TextLoader, CSVLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_cohere import CohereEmbeddings
    from langchain_qdrant import Qdrant
    from langchain_community.tools import DuckDuckGoSearchRun
    from langchain.docstore.document import Document

    # OCR imports with availability check
    try:
        import pytesseract
        from PIL import Image
        import pdf2image
        from docx import Document as DocxDocument
        OCR_AVAILABLE = True
    except ImportError:
        OCR_AVAILABLE = False

    # Add parent directory to path for imports
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    # Import centralized configuration
    from prompts_and_logic import RAGConfiguration, RAGAnalysisSteps, UIConfiguration, token_manager

    load_dotenv()

    # Add custom CSS for improved UI
    st.markdown("""
    <style>
        /* Pipeline indicator styling */
        .pipeline-step {
            display: inline-block;
            padding: 5px 10px;
            margin-right: 5px;
            border-radius: 15px;
            font-size: 0.8em;
        }
        .pipeline-step.active {
            background-color: #f0f2f6;
            border: 1px solid #4e8cff;
        }
        .pipeline-step.completed {
            background-color: #e6f3e6;
            border: 1px solid #4CAF50;
        }
        .pipeline-step.pending {
            background-color: #f5f5f5;
            border: 1px solid #ddd;
            color: #888;
        }
        
        /* File type styling */
        .file-type {
            display: inline-block;
            padding: 3px 8px;
            border-radius: 10px;
            font-size: 0.8em;
            margin-right: 5px;
            background-color: #f0f2f6;
        }
        
        /* Evaluation styling */
        .eval-dimension {
            margin-bottom: 10px;
        }
        .eval-score {
            font-weight: bold;
            color: #4e8cff;
        }
    </style>
    """, unsafe_allow_html=True)

    def get_api_keys():
        """Get API keys from environment or structured secrets."""
        # Cohere API Key (primary)
        cohere_key = os.getenv("COHERE_API_KEY")
        if not cohere_key and hasattr(st, "secrets"):
            if "COHERE_API_KEY" in st.secrets:
                cohere_key = st.secrets["COHERE_API_KEY"]
            elif "cohere" in st.secrets and "COHERE_API_KEY" in st.secrets.cohere:
                cohere_key = st.secrets.cohere["COHERE_API_KEY"]
        
        # OpenAI API Key (fallback)
        openai_key = os.getenv("OPENAI_API_KEY") 
        if not openai_key and hasattr(st, "secrets"):
            if "OPENAI_API_KEY" in st.secrets:
                openai_key = st.secrets["OPENAI_API_KEY"]
            elif "openai" in st.secrets and "OPENAI_API_KEY" in st.secrets.openai:
                openai_key = st.secrets.openai["OPENAI_API_KEY"]
        
        # Groq API Key (additional provider)
        groq_key = os.getenv("GROQ_API_KEY") 
        if not groq_key and hasattr(st, "secrets"):
            if "GROQ_API_KEY" in st.secrets:
                groq_key = st.secrets["GROQ_API_KEY"]
            elif "groq" in st.secrets and "GROQ_API_KEY" in st.secrets.groq:
                groq_key = st.secrets.groq["GROQ_API_KEY"]
        
        return cohere_key, openai_key, groq_key

    def create_gpt4all_embeddings():
        """Create GPT4All embeddings for local embedding generation."""
        try:
            from gpt4all import Embed4All
            
            # Initialize the GPT4All embedding model
            embedding_model = Embed4All(model_name=RAGConfiguration.GPT4ALL_EMBEDDING_MODEL)
            
            # Create a LangChain compatible embeddings class
            class GPT4AllEmbeddings:
                def __init__(self, model):
                    self.model = model
                    self.embedding_dimension = RAGConfiguration.GPT4ALL_EMBEDDING_DIMENSIONS
                
                def embed_documents(self, texts):
                    """Generate embeddings for a list of documents."""
                    return [self.model.embed(text) for text in texts]
                
                def embed_query(self, text):
                    """Generate embeddings for a query."""
                    return self.model.embed(text)
            
            return GPT4AllEmbeddings(embedding_model)
        except ImportError:
            st.error("GPT4All is not installed. Install with: pip install gpt4all")
            return None
        except Exception as e:
            st.error(f"Error initializing GPT4All embeddings: {e}")
            return None

    def load_compensation_database_with_fallback(cohere_key, openai_key=None):
        """Load the compensation database with automatic fallback to OpenAI if Cohere hits rate limits."""
        # Get the project root directory
        current_file = os.path.abspath(__file__)
        current_dir = os.path.dirname(current_file)
        project_root = os.path.dirname(current_dir)
        data_path = os.path.join(project_root, "data.csv")
        
        if not os.path.exists(data_path):
            return None, "Compensation database not found"
        
        try:
            df = pd.read_csv(data_path)
            
            # Convert to text documents for embedding
            documents = []
            for _, row in df.iterrows():
                # Format numbers with commas, handle 'N/A' gracefully
                def fmt(val):
                    if isinstance(val, (int, float)) and not pd.isna(val):
                        return f"{int(val):,}"
                    return str(val) if val is not None else "N/A"
                
                base_salary = row.get('base_salary_usd', 'N/A')
                bonus = row.get('bonus_usd', 'N/A')
                equity = row.get('equity_value_usd', 'N/A')
                total_comp = (
                    (base_salary if isinstance(base_salary, (int, float)) else 0) +
                    (bonus if isinstance(bonus, (int, float)) else 0) +
                    (equity if isinstance(equity, (int, float)) else 0)
                )
                
                doc_text = f"""
                Job Title: {row.get('job_title', 'N/A')}
                Level: {row.get('job_level', 'N/A')}
                Location: {row.get('location', 'N/A')}
                Base Salary: ${fmt(base_salary)} USD
                Bonus: ${fmt(bonus)} USD
                Equity Value: ${fmt(equity)} USD
                Company Stage: {row.get('company_stage', 'N/A')}
                Offer Outcome: {row.get('offer_outcome', 'N/A')}
                Candidate Preference: {row.get('candidate_preference', 'N/A')}
                Notes: {row.get('notes', 'N/A')}
                
                Total Compensation: ${fmt(total_comp)} USD
                Role Summary: {row.get('job_title', 'Engineer')} at {row.get('job_level', 'L4')} level in {row.get('location', 'Unknown')} earning ${fmt(base_salary)} base salary
                """
                documents.append(doc_text.strip())
            
            # Convert to LangChain documents
            docs = [Document(page_content=doc, metadata={"source": "compensation_database", "record_id": i}) 
                   for i, doc in enumerate(documents)]
            
            # Try Cohere embeddings first
            try:
                embeddings = CohereEmbeddings(
                    model="embed-english-v3.0",
                    cohere_api_key=cohere_key
                )
                
                # Test the embeddings with a small sample
                _ = embeddings.embed_query("test query")
                
                # If successful, create the full vectorstore
                vectorstore = Qdrant.from_documents(
                    docs, 
                    embeddings, 
                    location=":memory:",
                    collection_name="compensation_db"
                )
                
                return vectorstore, None
                
            except Exception as cohere_error:
                error_str = str(cohere_error)
                
                # Check if it's a rate limit error (status 429)
                if "429" in error_str or "rate limit" in error_str.lower() or "trial key" in error_str.lower():
                    if openai_key:
                        try:
                            from langchain_openai import OpenAIEmbeddings
                            
                            embeddings = OpenAIEmbeddings(
                                model="text-embedding-ada-002",
                                openai_api_key=openai_key
                            )
                            
                            # Create Qdrant vectorstore with OpenAI embeddings
                            vectorstore = Qdrant.from_documents(
                                docs, 
                                embeddings, 
                                location=":memory:",
                                collection_name="compensation_db"
                            )
                            
                            return vectorstore, None
                            
                        except Exception as openai_error:
                            return None, f"Both Cohere and OpenAI failed. Cohere: {error_str}, OpenAI: {str(openai_error)}"
                    else:
                        return None, f"Cohere rate limit reached and no OpenAI key available. Error: {error_str}"
                else:
                    # For other Cohere errors, don't fall back
                    return None, str(cohere_error)
        
        except Exception as e:
            return None, str(e)

    def process_uploaded_files(uploaded_files, cohere_key):
        """Process uploaded files and create FAISS vectorstore."""
        if not uploaded_files:
            return None, []
        
        all_docs = []
        processing_details = []
        
        with st.spinner('Processing uploaded files...'):
            for uploaded_file in uploaded_files:
                try:
                    # Check file size
                    file_size_mb = uploaded_file.size / (1024 * 1024)  # Convert to MB
                    if file_size_mb > RAGConfiguration.MAX_FILE_SIZE_MB:
                        st.warning(f"⚠️ File {uploaded_file.name} exceeds size limit of {RAGConfiguration.MAX_FILE_SIZE_MB}MB. Skipping.")
                        continue
                    
                    # Get file extension and check if supported
                    file_ext = uploaded_file.name.split('.')[-1].lower()
                    if file_ext not in RAGConfiguration.SUPPORTED_FILE_TYPES:
                        st.warning(f"⚠️ Unsupported file type: {file_ext}. Skipping {uploaded_file.name}")
                        continue
                    
                    # Show file processing indicator
                    file_type_info = RAGConfiguration.SUPPORTED_FILE_TYPES.get(file_ext, {})
                    file_icon = file_type_info.get('icon', '📄')
                    file_desc = file_type_info.get('description', 'Document')
                    
                    st.markdown(f"<div class='file-type'>{file_icon} {file_desc}</div> Processing {uploaded_file.name}...", unsafe_allow_html=True)
                    
                    # Save file temporarily
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        tmp_path = tmp_file.name
                    
                    # Process based on file type
                    if uploaded_file.type == "application/pdf":
                        loader = PyPDFLoader(tmp_path)
                        docs = loader.load()
                        processing_details.append(f"{file_icon} {uploaded_file.name}: PDF text extraction")
                        
                        # OCR for images in PDF if available
                        if OCR_AVAILABLE:
                            try:
                                images = pdf2image.convert_from_path(tmp_path)
                                for i, image in enumerate(images):
                                    text = pytesseract.image_to_string(image)
                                    if text.strip():
                                        docs.append(Document(
                                            page_content=text,
                                            metadata={"source": uploaded_file.name, "page": i+1, "type": "ocr"}
                                        ))
                                processing_details.append(f"🔍 {uploaded_file.name}: OCR text extraction from images")
                            except Exception as e:
                                processing_details.append(f"⚠️ {uploaded_file.name}: OCR failed - {str(e)}")
                        
                    elif uploaded_file.type == "text/plain":
                        loader = TextLoader(tmp_path)
                        docs = loader.load()
                        processing_details.append(f"{file_icon} {uploaded_file.name}: Plain text loaded")
                        
                    elif uploaded_file.type == "text/csv":
                        loader = CSVLoader(tmp_path)
                        docs = loader.load()
                        processing_details.append(f"{file_icon} {uploaded_file.name}: CSV data loaded")
                        
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and OCR_AVAILABLE:
                        doc = DocxDocument(tmp_path)
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        docs = [Document(page_content=text, metadata={"source": uploaded_file.name, "type": "docx"})]
                        processing_details.append(f"{file_icon} {uploaded_file.name}: Word document text extracted")
                        
                    elif OCR_AVAILABLE and uploaded_file.type.startswith('image/'):
                        image = Image.open(tmp_path)
                        text = pytesseract.image_to_string(image)
                        if text.strip():
                            docs = [Document(page_content=text, metadata={"source": uploaded_file.name, "type": "ocr"})]
                            processing_details.append(f"{file_icon} {uploaded_file.name}: OCR text extraction from image")
                        else:
                            processing_details.append(f"⚠️ {uploaded_file.name}: No text extracted from image")
                            docs = []
                        
                    else:
                        st.warning(f"Unsupported file type: {uploaded_file.type}")
                        continue
                    
                    # Add source metadata
                    for doc in docs:
                        doc.metadata['source_file'] = uploaded_file.name
                    
                    # Show preview of extracted text
                    if docs:
                        with st.expander(f"Preview: {uploaded_file.name}", expanded=False):
                            preview_text = docs[0].page_content[:200] + "..." if len(docs[0].page_content) > 200 else docs[0].page_content
                            st.text(preview_text)
                            st.caption(f"Extracted {len(docs)} document chunks")
                
                    all_docs.extend(docs)
                    
                    # Clean up
                    os.unlink(tmp_path)
                    
                except Exception as e:
                    st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        
        if all_docs:
            # Split documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=RAGConfiguration.CHUNK_SIZE,
                chunk_overlap=RAGConfiguration.CHUNK_OVERLAP
            )
            split_docs = text_splitter.split_documents(all_docs)
            
            # Choose embedding method based on configuration
            if RAGConfiguration.USE_LOCAL_EMBEDDINGS:
                # Use GPT4All for local embeddings
                show_search_analysis("Embedding", "🔮 Using GPT4All for local embeddings")
                embeddings = create_gpt4all_embeddings()
                if not embeddings:
                    st.warning("Falling back to Cohere embeddings due to GPT4All initialization failure")
                    embeddings = CohereEmbeddings(
                        model=RAGConfiguration.EMBEDDING_MODEL,
                        cohere_api_key=cohere_key
                    )
                processing_details.append("Using GPT4All local embeddings")
            else:
                # Use Cohere for cloud embeddings
                embeddings = CohereEmbeddings(
                    model=RAGConfiguration.EMBEDDING_MODEL,
                    cohere_api_key=cohere_key
                )
            
            # Create Qdrant vectorstore (in-memory for simplicity)
            vectorstore = Qdrant.from_documents(
                split_docs, 
                embeddings, 
                location=":memory:",
                collection_name="uploaded_docs"
            )
            
            return vectorstore, processing_details
        
        return None, processing_details

    def show_search_analysis(step, message, status="info"):
        """Display real-time search analysis"""
        timestamp = time.strftime("%H:%M:%S")
        
        # Store in session state for persistence
        if 'search_steps' not in st.session_state:
            st.session_state.search_steps = []
        
        st.session_state.search_steps.append({
            "timestamp": timestamp,
            "step": step,
            "message": message,
            "status": status
        })
        
        # Display in real-time
        if status == "success":
            st.success(f"**{step}**: {message}")
        elif status == "warning":
            st.warning(f"**{step}**: {message}")
        elif status == "error":
            st.error(f"**{step}**: {message}")
        else:
            st.info(f"**{step}**: {message}")

    def display_pipeline_progress(current_step):
        """Display a visual pipeline progress indicator"""
        pipeline_steps = UIConfiguration.PIPELINE_STEPS
        
        # Create HTML for pipeline steps
        html_parts = []
        for step in pipeline_steps:
            step_id = step["id"]
            step_name = step["name"]
            step_icon = step["icon"]
            
            if step_id == current_step:
                status_class = "active"
            elif pipeline_steps.index(step) < [s["id"] for s in pipeline_steps].index(current_step):
                status_class = "completed"
            else:
                status_class = "pending"
            
            html_parts.append(f'<div class="pipeline-step {status_class}">{step_icon} {step_name}</div>')
        
        # Display the pipeline progress
        st.markdown(
            f"""<div style="margin-bottom: 10px;">{''.join(html_parts)}</div>""", 
            unsafe_allow_html=True
        )

    def search_multi_source(query, db_vectorstore, uploads_vectorstore, cohere_client, openai_client=None, top_k=None):
        """Search across multiple sources with detailed analysis"""
        if top_k is None:
            top_k = RAGConfiguration.TOP_K_RESULTS
            
        all_results = []
        search_details = {
            "db_results": 0,
            "upload_results": 0,
            "total_results": 0,
            "sources_searched": [],
            "rerank_applied": False,
            "web_fallback": False,
            "current_step": "embedding"  # Track current pipeline step
        }
        
        # Display pipeline progress indicator
        display_pipeline_progress(search_details["current_step"])
        
        # Step 1: Search compensation database
        if db_vectorstore:
            show_search_analysis("Embedding", "🔮 Converting query to embeddings using embed-english-v3.0...")
            show_search_analysis("Database Search", "🔍 Searching compensation database...")
            search_details["current_step"] = "search"
            display_pipeline_progress(search_details["current_step"])
            
            try:
                db_docs = db_vectorstore.similarity_search(query, k=top_k)
                search_details["db_results"] = len(db_docs)
                search_details["sources_searched"].append("Compensation Database")
                
                for doc in db_docs:
                    doc.metadata["search_source"] = "database"
                all_results.extend(db_docs)
                
                show_search_analysis("Database Search", f"Found {len(db_docs)} relevant records in compensation database", "success")
            except Exception as e:
                show_search_analysis("Database Search", f"Error searching database: {str(e)}", "error")
        
        # Step 2: Search uploaded documents
        if uploads_vectorstore:
            if not db_vectorstore:  # Only show embedding step if not already shown
                show_search_analysis("Embedding", "🔮 Converting query to embeddings using embed-english-v3.0...")
            show_search_analysis("Document Search", "📄 Searching uploaded documents...")
            search_details["current_step"] = "search"
            display_pipeline_progress(search_details["current_step"])
            
            try:
                upload_docs = uploads_vectorstore.similarity_search(query, k=top_k)
                search_details["upload_results"] = len(upload_docs)
                search_details["sources_searched"].append("Uploaded Documents")
                
                for doc in upload_docs:
                    doc.metadata["search_source"] = "uploads"
                all_results.extend(upload_docs)
                
                show_search_analysis("Document Search", f"Found {len(upload_docs)} relevant chunks in uploaded documents", "success")
            except Exception as e:
                show_search_analysis("Document Search", f"Error searching uploads: {str(e)}", "error")
        
        # Step 3: Always rerank results if we have any results and AI client
        if len(all_results) > 0 and (cohere_client or openai_client) and RAGConfiguration.should_apply_reranking(all_results, cohere_client or openai_client):
            show_search_analysis("Reranking", "🎯 Reranking results using AI...")
            search_details["current_step"] = "reranking"
            display_pipeline_progress(search_details["current_step"])
            
            try:
                # Try Cohere reranking first
                if cohere_client:
                    doc_texts = [doc.page_content for doc in all_results]
                    rerank_response = cohere_client.rerank(
                        model=RAGConfiguration.RERANK_MODEL,
                        query=query,
                        documents=doc_texts,
                        top_n=min(top_k, len(all_results))
                    )
                    
                    # Process reranking results using centralized configuration
                    reranked_results, rerank_scores, low_relevance = RAGConfiguration.process_rerank_results(
                        rerank_response, 
                        all_results, 
                        RAGConfiguration.RERANK_MINIMUM_RELEVANCE_THRESHOLD
                    )
                    rerank_method = "Cohere"
                    
                # Fallback to OpenAI reranking if Cohere is not available
                elif openai_client:
                    reranked_results, rerank_scores, low_relevance = RAGConfiguration.openai_rerank_results(
                        query=query,
                        documents=all_results,
                        openai_client=openai_client,
                        top_n=min(top_k, len(all_results)),
                        min_threshold=RAGConfiguration.RERANK_MINIMUM_RELEVANCE_THRESHOLD
                    )
                    rerank_method = "OpenAI"
                
                search_details["rerank_applied"] = True
                search_details["total_results"] = len(reranked_results)
                search_details["low_relevance"] = low_relevance
                
                # Check if best result is actually relevant for fallback decision
                if reranked_results and low_relevance:
                    top_score = reranked_results[0].metadata.get("rerank_score", 0)
                    show_search_analysis(
                        "Relevance Check", 
                        f"Top result relevance ({top_score:.2f}) below threshold {RAGConfiguration.RERANK_RELEVANCE_THRESHOLD} - will try web fallback", 
                        "warning"
                    )
                
                show_search_analysis("Internal Reranking", f"Internal sources: {len(reranked_results)} relevant results using {rerank_method} (scores: {[f'{s:.2f}' for s in rerank_scores[:3]]})", "success" if len(reranked_results) > 0 else "warning")
                
                return reranked_results, search_details
                
            except Exception as e:
                show_search_analysis("Reranking", f"Reranking failed: {str(e)}, using similarity results", "warning")
        
        search_details["total_results"] = len(all_results)
        # Show message for 0 internal results to clarify web fallback will be tried
        if len(all_results) == 0:
            show_search_analysis("Internal Search", "0 relevant results from internal sources - will try web search", "warning")
        
        return all_results[:top_k], search_details

    def web_search_fallback(query, cohere_client):
        """Enhanced web search with detailed RAG pipeline analysis"""
        show_search_analysis("Web Fallback", "🌐 No relevant data found internally, searching the web...")
        
        # Update pipeline progress
        search_details = {"current_step": "search"}
        display_pipeline_progress(search_details["current_step"])
        
        try:
            search = DuckDuckGoSearchRun(num_results=RAGConfiguration.WEB_SEARCH_RESULTS)
            
            # Try multiple search strategies from configuration
            search_queries = RAGConfiguration.get_web_search_queries(query)
            
            best_results = None
            successful_query = None
            
            for search_query in search_queries:
                try:
                    show_search_analysis("Web Search", f"Trying search: '{search_query[:50]}...'")
                    results = search.run(search_query)
                    if results and len(results.strip()) > 100:
                        best_results = results
                        successful_query = search_query
                        break
                except Exception as e:
                    show_search_analysis("Web Search", f"Search attempt failed: {str(e)}", "warning")
                    continue
            
            if best_results and cohere_client:
                # Split web results for better processing
                web_chunks = [chunk.strip() for chunk in best_results.split('\n\n') if len(chunk.strip()) > RAGConfiguration.get_web_search_config()["chunk_min_length"]]
                
                # Always try to rerank if we have any chunks (even just 1)
                if len(web_chunks) >= 1:
                    try:
                        # Show web reranking step even for single chunks
                        show_search_analysis("Web Reranking", f"🎯 Reranking {len(web_chunks)} web chunks using Cohere...")
                        
                        # Update pipeline progress
                        search_details["current_step"] = "reranking"
                        display_pipeline_progress(search_details["current_step"])
                        
                        rerank_response = cohere_client.rerank(
                            model=RAGConfiguration.RERANK_MODEL,
                            query=query,
                            documents=web_chunks,
                            top_n=min(3, len(web_chunks))
                        )
                        
                        relevant_chunks = []
                        rerank_scores = []
                        for result in rerank_response.results:
                            if result.relevance_score > RAGConfiguration.WEB_SEARCH_THRESHOLD:
                                relevant_chunks.append(web_chunks[result.index])
                                rerank_scores.append(result.relevance_score)
                        
                        if relevant_chunks:
                            processed_results = "\n\n".join(relevant_chunks)
                            show_search_analysis("Web Reranking Results", f"✅ Selected {len(relevant_chunks)} relevant chunks (scores: {[f'{s:.2f}' for s in rerank_scores]})", "success")
                            show_search_analysis("Web Processing", "✅ Web RAG pipeline complete", "success")
                            return processed_results, ["🌐 DuckDuckGo + Cohere rerank"], True
                        else:
                            # All chunks below threshold, but show the scores anyway
                            all_scores = [result.relevance_score for result in rerank_response.results]
                            show_search_analysis("Web Reranking Results", f"⚠️ All chunks below threshold {RAGConfiguration.WEB_SEARCH_THRESHOLD} (scores: {[f'{s:.2f}' for s in all_scores]})", "warning")
                            # Fall back to using best result anyway
                            processed_results = web_chunks[rerank_response.results[0].index] if rerank_response.results else best_results
                            return processed_results, ["🌐 DuckDuckGo web search"], True
                            
                    except Exception as e:
                        show_search_analysis("Web Reranking", f"Reranking failed: {str(e)}", "warning")
                
                show_search_analysis("Web Search", f"Found web results for '{successful_query}' (no reranking applied)", "success")
                return best_results[:RAGConfiguration.get_web_search_config()["max_results_to_process"]], ["🌐 DuckDuckGo web search"], True
            
            elif best_results:
                show_search_analysis("Web Search", f"Found basic web results for '{successful_query}' (no Cohere client)", "success")
                return best_results, ["🌐 DuckDuckGo web search"], True
            
            else:
                show_search_analysis("Web Search", "No relevant web results found", "warning")
                return None, [], True
                
        except Exception as e:
            show_search_analysis("Web Search", f"Web search failed: {str(e)}", "error")
            return None, [], True

    def generate_rag_response(query, context_docs, sources, cohere_client, openai_client=None, groq_client=None):
        """Generate response using RAG context with detailed analysis and cross-source stitching"""
        show_search_analysis(*RAGAnalysisSteps.response_generation())
        
        # Update pipeline progress
        search_details = {"current_step": "generation"}
        display_pipeline_progress(search_details["current_step"])
        
        # Analyze and stitch information from multiple sources
        db_context = []
        upload_context = []
        web_context = []
        
        for doc in context_docs:
            source = doc.metadata.get("search_source", doc.metadata.get("source", "unknown"))
            if source == "database":
                db_context.append(doc.page_content)
            elif source == "uploads":
                upload_context.append(doc.page_content)
            elif source == "web_search":
                web_context.append(doc.page_content)
        
        # Show cross-source analysis
        source_count = len([x for x in [db_context, upload_context, web_context] if x])
        if source_count > 1:
            show_search_analysis(*RAGAnalysisSteps.cross_source_analysis(source_count))
        
        # Join contexts by source
        db_context_str = chr(10).join(db_context[:3]) if db_context else ""  # Limit to top 3 for space
        upload_context_str = chr(10).join(upload_context[:2]) if upload_context else ""  # Limit to top 2
        web_context_str = chr(10).join(web_context[:2]) if web_context else ""  # Limit to top 2
        
        # Get prompt from centralized configuration
        prompt = RAGConfiguration.get_cross_source_prompt(
            query=query,
            db_context=db_context_str,
            upload_context=upload_context_str,
            web_context=web_context_str
        )
        
        # Try Cohere first
        if cohere_client:
            try:
                show_search_analysis(*RAGAnalysisSteps.ai_generation("Cohere command-r-plus"))
                response = cohere_client.chat(
                    model=RAGConfiguration.GENERATION_MODEL_PRIMARY,
                    message=prompt,
                    temperature=RAGConfiguration.GENERATION_TEMPERATURE,
                    max_tokens=RAGConfiguration.GENERATION_MAX_TOKENS
                )
                show_search_analysis(*RAGAnalysisSteps.response_success())
                return response.text, "Cohere"
            except Exception as e:
                show_search_analysis("AI Generation", f"Cohere failed: {str(e)}, trying Groq...", "warning")
        
        # Try Groq as second option
        if groq_client:
            try:
                show_search_analysis(*RAGAnalysisSteps.ai_generation(f"Groq {RAGConfiguration.GENERATION_MODEL_GROQ_PRIMARY}"))
                # Use Groq API for chat completions with primary model
                response = groq_client.chat.completions.create(
                    model=RAGConfiguration.GENERATION_MODEL_GROQ_PRIMARY,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=RAGConfiguration.GENERATION_TEMPERATURE,
                    max_tokens=RAGConfiguration.GENERATION_MAX_TOKENS
                )
                show_search_analysis(*RAGAnalysisSteps.response_success())
                return response.choices[0].message.content, "Groq (Llama 3 70B)"
            except Exception as e:
                # Try fallback to secondary Groq model
                try:
                    show_search_analysis("AI Generation", f"Primary Groq model failed: {str(e)}, trying secondary Groq model...", "warning")
                    show_search_analysis(*RAGAnalysisSteps.ai_generation(f"Groq {RAGConfiguration.GENERATION_MODEL_GROQ_SECONDARY}"))
                    
                    response = groq_client.chat.completions.create(
                        model=RAGConfiguration.GENERATION_MODEL_GROQ_SECONDARY,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=RAGConfiguration.GENERATION_TEMPERATURE,
                        max_tokens=RAGConfiguration.GENERATION_MAX_TOKENS
                    )
                    show_search_analysis(*RAGAnalysisSteps.response_success())
                    return response.choices[0].message.content, "Groq (Llama 3 8B)"
                except Exception as inner_e:
                    show_search_analysis("AI Generation", f"Secondary Groq model also failed: {str(inner_e)}, trying OpenAI...", "warning")
    
        # Fallback to OpenAI
        if openai_client:
            try:
                show_search_analysis(*RAGAnalysisSteps.ai_generation("OpenAI GPT-3.5-turbo"))
                response = openai_client.chat.completions.create(
                    model=RAGConfiguration.GENERATION_MODEL_OPENAI,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=RAGConfiguration.GENERATION_TEMPERATURE,
                    max_tokens=RAGConfiguration.GENERATION_FALLBACK_MAX_TOKENS
                )
                show_search_analysis(*RAGAnalysisSteps.response_success())
                return response.choices[0].message.content, "OpenAI"
            except Exception as e:
                show_search_analysis("AI Generation", f"OpenAI failed: {str(e)}", "error")
                return f"Error generating response: {e}", "Error"
    
        return "No AI providers available. Please configure API keys.", "Error"

    def display_search_analysis():
        """Display the search analysis in an expandable section"""
        if 'search_steps' in st.session_state and st.session_state.search_steps:
            with st.expander(UIConfiguration.SEARCH_ANALYSIS_HEADER, expanded=False):
                st.markdown(f"### {UIConfiguration.SEARCH_ANALYSIS_SUBHEADER}")
                
                # Group steps by category for better organization
                step_categories = {}
                
                for step_info in st.session_state.search_steps:
                    step_name = step_info["step"]
                    category = step_name.split()[0] if " " in step_name else step_name
                    
                    if category not in step_categories:
                        step_categories[category] = []
                        
                    step_categories[category].append(step_info)
                
                # Display steps by category with better formatting
                for category, steps in step_categories.items():
                    with st.container():
                        st.markdown(f"**{category}**")
                        
                        for step_info in steps:
                            status_icon = UIConfiguration.STATUS_ICONS.get(step_info["status"], "ℹ️")
                            
                            # Format timestamp and message
                            timestamp = step_info["timestamp"]
                            message = step_info["message"]
                            
                            # Use different styling based on status
                            if step_info["status"] == "success":
                                st.markdown(f"{status_icon} `{timestamp}` {message}")
                            elif step_info["status"] == "warning":
                                st.warning(f"`{timestamp}` {message}")
                            elif step_info["status"] == "error":
                                st.error(f"`{timestamp}` {message}")
                            else:
                                st.info(f"`{timestamp}` {message}")
                    
                    st.markdown("---")

    def simple_rag_evaluate(query, context_docs, response, cohere_client=None, openai_client=None, top_rerank_score=None):
        """Simplified RAG evaluation function"""
        eval_id = str(uuid.uuid4())
        
        # Use either Cohere or OpenAI for evaluation
        if cohere_client:
            try:
                # Join contexts into a single string
                contexts_str = "\n\n".join([doc.page_content for doc in context_docs])
                
                # Get the evaluation prompt from centralized configuration
                prompt = RAGConfiguration.get_simple_rag_evaluation_prompt(query, contexts_str, response, top_rerank_score=top_rerank_score)
                
                # Call Cohere API
                eval_response = cohere_client.chat(
                    message=prompt,
                    model=RAGConfiguration.GENERATION_MODEL_PRIMARY,
                    temperature=0.1
                )
                
                # Parse the JSON response
                try:
                    json_start = eval_response.text.find('{')
                    json_end = eval_response.text.rfind('}') + 1
                    json_str = eval_response.text[json_start:json_end]
                    evaluation = json.loads(json_str)
                    return evaluation, eval_id
                except Exception as e:
                    st.warning(f"Error parsing evaluation response: {str(e)}")
                    return None, eval_id
                
            except Exception as e:
                st.warning(f"Evaluation error: {str(e)}")
                return None, eval_id
                
        elif openai_client:
            try:
                # Join contexts into a single string
                contexts_str = "\n\n".join([doc.page_content for doc in context_docs])
                
                # Get the evaluation prompt from centralized configuration
                prompt = RAGConfiguration.get_simple_rag_evaluation_prompt(query, contexts_str, response, top_rerank_score=top_rerank_score)
                
                # Call OpenAI API
                eval_response = openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are an expert RAG system evaluator. Provide your evaluation in valid JSON format only."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1
                )
                
                # Parse the JSON response
                try:
                    json_str = eval_response.choices[0].message.content
                    json_start = json_str.find('{')
                    json_end = json_str.rfind('}') + 1
                    json_str = json_str[json_start:json_end]
                    evaluation = json.loads(json_str)
                    return evaluation, eval_id
                except Exception as e:
                    st.warning(f"Error parsing evaluation response: {str(e)}")
                    return None, eval_id
                    
            except Exception as e:
                st.warning(f"Evaluation error: {str(e)}")
                return None, eval_id
        
        return None, eval_id

    def display_evaluation_results(evaluation):
        """Display the simplified evaluation results in a clean format"""
        if not evaluation:
            st.warning("Evaluation failed or returned no results.")
            return
        
        # Get overall score
        overall_score = evaluation.get("overall_score", 0)
        
        # Display score with gauge chart or simple metric
        col1, col2 = st.columns([1, 3])
        with col1:
            st.metric("Overall", f"{overall_score:.1f}/10")
        
        # Display dimension scores
        with col2:
            dim_scores = []
            for dim in ["relevance", "factual_accuracy", "groundedness"]:
                if dim in evaluation:
                    score = evaluation[dim]["score"]
                    dim_scores.append((dim.replace("_", " ").title(), score))
            
            # Create horizontal bar chart for dimension scores
            if dim_scores:
                chart_data = pd.DataFrame({
                    "Dimension": [d[0] for d in dim_scores],
                    "Score": [d[1] for d in dim_scores]
                })
                st.bar_chart(chart_data, x="Dimension", y="Score", height=100)
        
        # Display strengths and areas for improvement
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Strengths")
            strengths = evaluation.get("strengths", [])
            for strength in strengths:
                st.markdown(f"✅ {strength}")
        
        with col2:
            st.subheader("Areas for Improvement")
            areas = evaluation.get("areas_for_improvement", [])
            for area in areas:
                st.markdown(f"🔍 {area}")

    def main():
        # Use the better welcome message from UIConfiguration at the top
        st.markdown(f"### {UIConfiguration.WELCOME_MESSAGE}")
        
        # Get API keys
        cohere_key, openai_key, groq_key = get_api_keys()
        gemini_key = None
        if hasattr(st, "secrets") and "gemini" in st.secrets and "GEMINI_API_KEY" in st.secrets.gemini:
            gemini_key = st.secrets.gemini["GEMINI_API_KEY"]

        # Main provider: OpenAI
        openai_client = None
        if openai_key:
            openai_client = OpenAI(api_key=openai_key)

        # Fallback: Gemini
        gemini_client = None
        if not openai_client and gemini_key:
            try:
                from google.generativeai import GenerativeModel
                gemini_client = GenerativeModel(model_name="gemini-pro", api_key=gemini_key)
            except ImportError:
                st.warning("Google Generative AI SDK not installed. Install with: pip install google-generativeai")

        # Fallback: Cohere
        cohere_client = None
        if not openai_client and not gemini_client and cohere_key:
            cohere_client = cohere.Client(cohere_key)
            os.environ["CHROMA_COHERE_API_KEY"] = cohere_key

        # Initialize session state
        if 'rag_messages' not in st.session_state:
            st.session_state.rag_messages = []
        
        if 'db_vectorstore' not in st.session_state:
            st.session_state.db_vectorstore = None
        
        if 'uploads_vectorstore' not in st.session_state:
            st.session_state.uploads_vectorstore = None
        
        # Add session state for message context (for evaluations)
        if 'message_context' not in st.session_state:
            st.session_state.message_context = {}
        
        # Initialize evaluations state
        if 'evaluations' not in st.session_state:
            st.session_state.evaluations = {}
        
        # Automatically load the compensation database if not already loaded
        if st.session_state.db_vectorstore is None:
            with st.spinner("Loading internal database..."):
                # Use the new fallback function that handles Cohere rate limits
                db_vectorstore, error = load_compensation_database_with_fallback(cohere_key, openai_key)
                if error:
                    if "429" in str(error) or "rate limit" in str(error).lower():
                        st.warning(f"⚠️ Cohere API rate limit reached. {error}")
                        if openai_key:
                            st.info("✅ Automatically switched to OpenAI embeddings as backup!")
                        else:
                            st.error("❌ No OpenAI backup available. Please add OPENAI_API_KEY to continue.")
                    else:
                        st.error(f"❌ Error loading database: {error}")
                else:
                    st.session_state.db_vectorstore = db_vectorstore
                    st.success("✅ Internal database loaded automatically!")

        # Add simple reference questions
        if st.session_state.db_vectorstore is not None:
            st.markdown("**💡 Try asking:** *What's the average salary for L5 engineers?* • *Compare salaries between New York and London* • *Which roles offer the highest equity?*")
        
        # Sidebar for file upload and settings
        with st.sidebar:
            # Display token usage widget at the top of sidebar
            token_manager.display_usage_widget()
            
            # Move Data Sources to the top
            st.header(UIConfiguration.SIDEBAR_DATA_HEADER)
            
            # Show current data sources at the top
            sources_status = []
            if st.session_state.db_vectorstore:
                sources_status.append("✅ Compensation Database")
            else:
                sources_status.append("❌ Compensation Database")
                
            uploaded_files = st.file_uploader(
                UIConfiguration.SIDEBAR_UPLOAD_LABEL,
                accept_multiple_files=True,
                type=list(RAGConfiguration.SUPPORTED_FILE_TYPES.keys()),
                help=UIConfiguration.SIDEBAR_UPLOAD_HELP
            )
            
            if st.session_state.uploads_vectorstore:
                sources_status.append(f"✅ Uploaded Documents ({len(uploaded_files) if uploaded_files else 0} files)")
            else:
                sources_status.append("❌ No uploaded documents")
            
            st.markdown("**Current Sources:**")
            for status in sources_status:
                st.markdown(f"• {status}")
            
            # Keep the manual load button for refreshing if needed
            if st.button(UIConfiguration.RELOAD_DB_BUTTON):
                with st.spinner("Reloading compensation database..."):
                    db_vectorstore, error = load_compensation_database_with_fallback(cohere_key, openai_key)
                    if error:
                        st.error(f"❌ Error loading database: {error}")
                    else:
                        st.session_state.db_vectorstore = db_vectorstore
                        st.success("✅ Compensation database reloaded!")
            
            st.markdown("---")
            st.header(UIConfiguration.SIDEBAR_UPLOAD_HEADER)
            
            # Create a styled upload area
            st.markdown("""
            <div class="upload-area">
                <div style="font-size: 2em;">📄</div>
                <div>Drag and drop files here</div>
                <div style="font-size: 0.8em; color: #666;">Supported formats: PDF, CSV, TXT, DOCX, PNG, JPG</div>
            </div>
            """, unsafe_allow_html=True)
            
            # Display supported file types with icons
            st.markdown("### Supported File Types")
            file_types_html = ""
            for ext, info in RAGConfiguration.SUPPORTED_FILE_TYPES.items():
                file_types_html += f'<div class="file-type">{info["icon"]} {ext.upper()}</div> '
            st.markdown(f"<div>{file_types_html}</div>", unsafe_allow_html=True)
            
            if uploaded_files:
                # Show file size information
                total_size_mb = sum(file.size for file in uploaded_files) / (1024 * 1024)
                st.caption(f"Total size: {total_size_mb:.2f}MB / {RAGConfiguration.MAX_FILE_SIZE_MB}MB limit per file")
                
                if st.button(UIConfiguration.PROCESS_FILES_BUTTON):
                    uploads_vectorstore, processing_details = process_uploaded_files(uploaded_files, cohere_key)
                    
                    if uploads_vectorstore:
                        st.session_state.uploads_vectorstore = uploads_vectorstore
                        st.success(f"✅ Processed {len(uploaded_files)} files successfully!")
                        
                        with st.expander("📋 Processing Details"):
                            for detail in processing_details:
                                st.write(f"• {detail}")
                    else:
                        st.error("❌ Failed to process uploaded files")
            
            # Settings
            st.markdown("---")
            st.header(UIConfiguration.SIDEBAR_SETTINGS_HEADER)
            
            web_fallback = st.checkbox(
                UIConfiguration.WEB_FALLBACK_TOGGLE, 
                value=True, 
                help=UIConfiguration.WEB_FALLBACK_HELP
            )
            
            # Add GPT4All local embeddings toggle
            use_local_embeddings = st.checkbox(
                "Use GPT4All Local Embeddings", 
                value=RAGConfiguration.USE_LOCAL_EMBEDDINGS,
                help="Use local GPT4All embeddings instead of cloud API calls (saves API tokens)"
            )
            
            # Update the configuration based on the checkbox
            if use_local_embeddings != RAGConfiguration.USE_LOCAL_EMBEDDINGS:
                RAGConfiguration.USE_LOCAL_EMBEDDINGS = use_local_embeddings
                st.info(f"{'Enabled' if use_local_embeddings else 'Disabled'} local embeddings.")
            
            if st.button(UIConfiguration.CLEAR_CHAT_BUTTON):
                st.session_state.rag_messages = []
                if 'search_steps' in st.session_state:
                    st.session_state.search_steps = []
                if 'evaluations' in st.session_state:
                    st.session_state.evaluations = {}
                if 'message_context' in st.session_state:
                    st.session_state.message_context = {}
                st.success("Chat history cleared!")
                st.rerun()
        
        # Display chat messages
        for i, message in enumerate(st.session_state.rag_messages):
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
                # Add evaluation button after assistant messages
                if message["role"] == "assistant" and i > 0:
                    # Message ID for tracking evaluations
                    msg_id = f"msg_{i}"
                    
                    # Show existing evaluation if available
                    if msg_id in st.session_state.evaluations:
                        evaluation = st.session_state.evaluations[msg_id]
                        display_evaluation_results(evaluation)
                    elif st.button(f"🔍 Evaluate Response", key=f"eval_btn_{i}"):
                        # Get the query and context docs
                        query = st.session_state.rag_messages[i-1]["content"] if i > 0 else ""
                        context_docs = st.session_state.message_context.get(msg_id, [])
                        response = message["content"]
                        
                        # Get top rerank score if available
                        top_rerank_score = None
                        if context_docs and hasattr(context_docs[0], 'metadata') and 'rerank_score' in context_docs[0].metadata:
                            top_rerank_score = context_docs[0].metadata.get('rerank_score')
                        
                        # Run evaluation
                        with st.spinner("Evaluating response quality..."):
                            evaluation, eval_id = simple_rag_evaluate(
                                query=query, 
                                context_docs=context_docs, 
                                response=response,
                                cohere_client=cohere_client,
                                openai_client=openai_client,
                                top_rerank_score=top_rerank_score
                            )
                            
                            # Store the evaluation result
                            if evaluation:
                                st.session_state.evaluations[msg_id] = evaluation
                                st.rerun()
        
        # Chat input
        if prompt := st.chat_input(UIConfiguration.CHAT_INPUT_PLACEHOLDER):
            # Add user message to chat history
            st.session_state.rag_messages.append({"role": "user", "content": prompt})
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Display assistant response in chat message container
            with st.chat_message("assistant"):
                with st.spinner("Searching and generating response..."):
                    # Reset search steps for this new query
                    st.session_state.search_steps = []
                    
                    # Search internal sources first
                    db_vectorstore = st.session_state.db_vectorstore
                    uploads_vectorstore = st.session_state.uploads_vectorstore
                    
                    if not db_vectorstore and not uploads_vectorstore:
                        response_text = "Please load the compensation database or upload documents to chat with your data."
                        provider = "Error"
                        context_docs = []
                        sources = []
                    else:
                        context_docs, search_details = search_multi_source(
                            query=prompt,
                            db_vectorstore=db_vectorstore,
                            uploads_vectorstore=uploads_vectorstore,
                            cohere_client=cohere_client,
                            openai_client=openai_client
                        )
                        
                        # Check if web fallback is needed (low relevance or no results)
                        need_web_fallback = web_fallback and (
                            (search_details.get("total_results", 0) == 0) or
                            (search_details.get("low_relevance", False) and search_details.get("total_results", 0) < 2)
                        )
                        
                        # If no results or low relevance and web fallback enabled, search the web
                        web_results = None
                        if need_web_fallback:
                            web_results, web_sources, _ = web_search_fallback(prompt, cohere_client)
                            if web_results:
                                # Convert web results to document format
                                web_docs = [Document(page_content=web_results, metadata={"source": "web_search"})]
                                context_docs.extend(web_docs)
                        
                        # Display search analysis if available
                        display_search_analysis()
                        
                        # Generate RAG response if we have context
                        if context_docs:
                            response_text, provider = generate_rag_response(
                                query=prompt,
                                context_docs=context_docs,
                                sources=search_details.get("sources_searched", []),
                                cohere_client=cohere_client,
                                openai_client=openai_client,
                                groq_client=groq_client
                            )
                        else:
                            response_text = UIConfiguration.NO_RESULTS_ERROR
                            provider = "Error"
                    
                    # Store context docs for this message (for evaluation)
                    msg_id = f"msg_{len(st.session_state.rag_messages)}"
                    st.session_state.message_context[msg_id] = context_docs
                    
                    # Write the response
                    st.markdown(response_text)
                    
                    # Add assistant response to chat history
                    st.session_state.rag_messages.append({
                        "role": "assistant", 
                        "content": response_text,
                        "provider": provider
                    })
except Exception as e:
    st.error(f"App failed to load: {e}")
    st.exception(e)
    st.text("Traceback:")
    st.text(traceback.format_exc())

main()
